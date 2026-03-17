#!/usr/bin/env python
"""
Generate comprehensive Word documentation for the Graph Theory Analysis Pipeline.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime


def add_styled_paragraph(doc, text, style='Normal', bold=False, italic=False,
                         font_size=None, color=None, space_after=None, space_before=None):
    """Helper to add a styled paragraph."""
    p = doc.add_paragraph(text, style=style)
    if bold or italic or font_size or color:
        for run in p.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if font_size:
                run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_code_block(doc, code_text):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def build_document():
    doc = Document()

    # =========================================================================
    # STYLES
    # =========================================================================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Graph Theory Analysis Pipeline\nfor XCP-D Preprocessed fMRI Data')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Documentation & Scientific Reference')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph('')
    doc.add_paragraph('')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Pipeline Script: run_gt.py')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('')

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f'Document Generated: {datetime.date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (placeholder)
    # =========================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Pipeline Overview & Architecture',
        '3. Input Data: XCP-D Preprocessed Outputs',
        '4. Step 1 — Data Discovery & File Identification',
        '5. Step 2 — Loading Parcellated Time Series',
        '6. Step 3 — Functional Connectivity Matrix Construction',
        '7. Step 4 — Edge Handling Strategy',
        '8. Step 5 — Network Thresholding',
        '9. Step 6 — Graph Theory Metric Computation',
        '10. Step 7 — Small-World Analysis',
        '11. Step 8 — Threshold Sweep & AUC Integration',
        '12. Step 9 — Output Structure & Data Products',
        '13. Step 10 — Visualization',
        '14. Configuration Parameters & Defaults',
        '15. Supported Parcellations',
        '16. Mathematical Formulations',
        '17. Software Dependencies & Versions',
        '18. Usage Instructions',
        '19. Limitations & Considerations',
        '20. References',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'This document provides complete technical documentation for the graph theory analysis '
        'pipeline implemented in run_gt.py. The pipeline takes XCP-D preprocessed resting-state '
        'fMRI data (parcellated time series in CIFTI ptseries format) and computes a comprehensive '
        'suite of graph-theoretic metrics characterizing the topological organization of functional '
        'brain networks.'
    )

    doc.add_paragraph(
        'The pipeline is designed for batch processing of large neuroimaging datasets with multiple '
        'subjects, sessions, runs, and parcellation schemes. It produces both weighted and binary '
        'network metrics, performs small-world analysis with appropriate null models, conducts '
        'multi-threshold sensitivity analysis with area-under-the-curve (AUC) integration, and '
        'generates publication-quality visualizations.'
    )

    p = doc.add_paragraph('Key design principles:', style='List Bullet')
    principles = [
        'Applies Fisher z-transformation (z = arctanh(r)) to correlation matrices before graph analysis, '
        'which is essential for valid group-level averaging and statistical inference across the n ≈ 1,000 '
        'subject cohort. The z-transform normalizes the sampling distribution and stabilizes variance across edges.',
        'Uses positive-only edge handling by default, discarding negative correlations whose neurobiological '
        'interpretation in graph-theoretic frameworks remains debated.',
        'Employs proportional thresholding to ensure matched network density across subjects, enabling '
        'valid between-subject comparisons.',
        'Computes small-worldness exclusively on binarized networks using degree-preserving null models '
        '(Maslov–Sneppen rewiring), which is the mathematically well-defined approach.',
        'Performs threshold sweep with AUC integration to mitigate dependence on any single arbitrary threshold.',
        'Uses the Louvain algorithm with best-of-10-runs selection for modularity to address the stochastic '
        'nature of community detection.',
        'Per-subject output organization facilitating downstream group-level analyses.',
    ]
    for pr in principles[1:]:
        doc.add_paragraph(pr, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 2. PIPELINE OVERVIEW & ARCHITECTURE
    # =========================================================================
    doc.add_heading('2. Pipeline Overview & Architecture', level=1)

    doc.add_paragraph(
        'The pipeline follows a sequential processing architecture organized around the '
        'XCPGraphTheoryPipeline class. Processing proceeds through the following major stages:'
    )

    stages = [
        ('Data Discovery', 'Automated identification of subjects, sessions, runs, and available '
         'parcellation schemes from the XCP-D output directory structure.'),
        ('Time Series Loading', 'Loading of CIFTI ptseries.nii files with robust orientation '
         'detection and validation against expected parcel counts.'),
        ('Connectivity Estimation', 'Computation of Pearson correlation matrices from parcellated '
         'time series, yielding subject-level functional connectivity matrices in r-value space.'),
        ('Edge Preparation', 'Application of edge handling strategy (positive-only, absolute, or signed) '
         'to prepare the connectivity matrix for graph-theoretic analysis.'),
        ('Thresholding', 'Proportional (or absolute) thresholding across a range of density levels '
         'to construct sparse network representations.'),
        ('Metric Computation', 'Extraction of weighted and binary graph metrics at each threshold, '
         'including efficiency, clustering, modularity, centrality, and more.'),
        ('Small-World Analysis', 'Computation of small-worldness (σ) using degree-preserving randomized '
         'null models on binarized networks.'),
        ('AUC Integration', 'Numerical integration of metrics across thresholds using the trapezoidal '
         'rule, yielding threshold-independent summary measures.'),
        ('Output Generation', 'Saving of per-subject metrics CSVs, connectivity matrices, threshold '
         'sweep data, and publication-quality figures.'),
    ]

    for i, (name, desc) in enumerate(stages, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Stage {i}: {name}. ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        'The pipeline iterates over parcellations as the outer loop, subjects within each parcellation, '
        'sessions within each subject, and individual runs within each session. This structure allows '
        'flexible processing of longitudinal, multi-run datasets.'
    )

    doc.add_page_break()

    # =========================================================================
    # 3. INPUT DATA: XCP-D PREPROCESSED OUTPUTS
    # =========================================================================
    doc.add_heading('3. Input Data: XCP-D Preprocessed Outputs', level=1)

    doc.add_paragraph(
        'The pipeline expects input data from XCP-D (https://xcp-d.readthedocs.io/), a post-processing '
        'tool for fMRI data that has been minimally preprocessed by fMRIPrep or the HCP pipeline. XCP-D '
        'performs confound regression, bandpass filtering, and parcellation of the preprocessed BOLD data.'
    )

    doc.add_heading('3.1 Expected Directory Structure', level=2)
    doc.add_paragraph(
        'The XCP-D output directory is expected to follow BIDS-derivative conventions:'
    )
    add_code_block(doc,
        'xcp_d_output/\n'
        '├── sub-001/\n'
        '│   ├── ses-01/\n'
        '│   │   └── func/\n'
        '│   │       ├── sub-001_ses-01_task-rest_..._seg-Glasser_stat-mean_timeseries.ptseries.nii\n'
        '│   │       ├── sub-001_ses-01_task-rest_run-01_..._seg-Glasser_stat-mean_timeseries.ptseries.nii\n'
        '│   │       └── ...\n'
        '│   └── ses-02/\n'
        '│       └── func/\n'
        '└── sub-002/\n'
        '    └── ...'
    )

    doc.add_heading('3.2 File Format: CIFTI ptseries', level=2)
    doc.add_paragraph(
        'Input files are in CIFTI-2 ptseries.nii format (parcellated time series). These files contain '
        'a 2D matrix with dimensions (timepoints × parcels), where each column represents the mean BOLD '
        'signal within a cortical or subcortical parcel over time. The pipeline uses nibabel to load '
        'these files and includes robust orientation detection logic.'
    )

    doc.add_heading('3.3 Concatenated vs. Per-Run Files', level=2)
    doc.add_paragraph(
        'XCP-D may produce both concatenated files (all runs merged) and per-run files. The pipeline '
        'supports both modes via the USE_CONCATENATED configuration flag. When set to False (the default), '
        'the pipeline processes each run independently, enabling run-level quality assessment and the '
        'computation of session-level averages. Concatenated files are identified by the absence of a '
        '"run-XX" entity in the filename.'
    )

    doc.add_paragraph(
        'Scientific rationale: Processing per-run files preserves the ability to assess within-session '
        'variability and to exclude individual problematic runs. Session-level averages can be computed '
        'downstream, providing a more principled approach than concatenation which can introduce '
        'discontinuities at run boundaries.'
    )

    doc.add_page_break()

    # =========================================================================
    # 4. STEP 1 — DATA DISCOVERY
    # =========================================================================
    doc.add_heading('4. Step 1 — Data Discovery & File Identification', level=1)

    doc.add_paragraph(
        'The pipeline begins by automatically discovering all available data in the XCP-D output directory. '
        'This is handled by three functions:'
    )

    p = doc.add_paragraph()
    run = p.add_run('find_xcp_subjects(): ')
    run.bold = True
    p.add_run('Identifies all subject directories matching the "sub-*" pattern.')

    p = doc.add_paragraph()
    run = p.add_run('find_sessions(): ')
    run.bold = True
    p.add_run('For each subject, identifies session directories matching "ses-*".')

    p = doc.add_paragraph()
    run = p.add_run('find_ptseries_files(): ')
    run.bold = True
    p.add_run(
        'Within each func/ directory, locates ptseries files matching the expected naming convention. '
        'Files are filtered by task (default: "rest"), run type (concatenated vs. per-run), and '
        'parcellation scheme (via the "seg-{parcellation}" entity).'
    )

    doc.add_paragraph(
        'The get_available_parcellations() function surveys the first 5 subjects to determine which '
        'parcellation schemes are available in the dataset. Users can then select specific parcellations '
        'via command-line arguments or process all available schemes.'
    )

    doc.add_page_break()

    # =========================================================================
    # 5. STEP 2 — LOADING PARCELLATED TIME SERIES
    # =========================================================================
    doc.add_heading('5. Step 2 — Loading Parcellated Time Series', level=1)

    doc.add_paragraph(
        'The load_ptseries() function loads CIFTI ptseries files and returns a matrix with shape '
        '(timepoints × parcels). This step includes critical validation and orientation handling.'
    )

    doc.add_heading('5.1 Orientation Detection', level=2)
    doc.add_paragraph(
        'CIFTI ptseries files follow the convention of (timepoints × parcels). However, the pipeline '
        'implements a three-tier orientation detection strategy to handle potential inconsistencies:'
    )

    steps = [
        ('Primary check (known parcellations)', 'If the parcellation is in the EXPECTED_PARCELS dictionary, '
         'the pipeline checks whether dimension 1 (columns) matches the expected parcel count. If dimension 0 '
         '(rows) matches instead, the data is transposed.'),
        ('Ambiguity warning', 'If neither dimension matches the expected parcel count, a warning is issued '
         'and the pipeline falls back to the heuristic.'),
        ('Heuristic fallback', 'For unknown parcellations, the pipeline assumes that the larger dimension '
         'represents timepoints (since typical fMRI acquisitions have more timepoints than parcels). If '
         'shape[0] < shape[1], the data is transposed.'),
    ]
    for name, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('5.2 Data Validation', level=2)
    doc.add_paragraph(
        'After loading, the pipeline validates the data:'
    )
    validations = [
        'Rejects 1D data arrays (corrupt or improperly formatted files).',
        'Verifies parcel count against the EXPECTED_PARCELS dictionary for known parcellations.',
        'Rejects time series with fewer than 50 timepoints, as these provide insufficient data '
        'for reliable correlation estimation.',
    ]
    for v in validations:
        doc.add_paragraph(v, style='List Bullet')

    doc.add_paragraph(
        'Scientific rationale for the 50-timepoint minimum: Reliable estimation of Pearson correlation '
        'requires a sufficient number of observations. With typical TRs of 0.8–2.0 seconds, 50 timepoints '
        'corresponds to approximately 40–100 seconds of data, which is considered a minimum for stable '
        'functional connectivity estimation (Van Dijk et al., 2010).'
    )

    doc.add_page_break()

    # =========================================================================
    # 6. STEP 3 — FUNCTIONAL CONNECTIVITY MATRIX
    # =========================================================================
    doc.add_heading('6. Step 3 — Functional Connectivity Matrix Construction', level=1)

    doc.add_heading('6.1 Correlation Method', level=2)
    doc.add_paragraph(
        'The pipeline computes a functional connectivity matrix using Pearson correlation. For each pair '
        'of parcels (i, j), the Pearson correlation coefficient r_ij is computed between their respective '
        'time series. This yields a symmetric N×N matrix where N is the number of parcels.'
    )

    doc.add_paragraph(
        'The implementation uses numpy.corrcoef(), which computes the full pairwise correlation matrix '
        'efficiently. An alternative partial correlation method (via nilearn ConnectivityMeasure) is '
        'also available but is not the default.'
    )

    doc.add_heading('6.2 Fisher Z-Transformation', level=2)
    doc.add_paragraph(
        'After computing the Pearson correlation matrix, the pipeline applies the Fisher z-transformation '
        '(z = arctanh(r)) to all edge weights. This is a critical step that converts bounded correlation '
        'coefficients (r ∈ [−1, 1]) into unbounded, approximately normally distributed z-scores.'
    )

    doc.add_paragraph(
        'Implementation detail: Before applying arctanh, correlation values are clipped to the range '
        '[−0.999999, 0.999999] to avoid numerical infinities at r = ±1 (which only occur on the diagonal '
        'or in degenerate cases). The diagonal is then zeroed.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Scientific rationale: ')
    run.bold = True
    p.add_run(
        'The Fisher z-transformation is essential in this pipeline because the analysis is designed for '
        'group-level inference across a large cohort (n ≈ 1,000 subjects). There are several reasons '
        'why the z-transform is required in this context:'
    )

    z_reasons = [
        ('Normalizing the sampling distribution',
         'The sampling distribution of Pearson\'s r is skewed, particularly for large |r| values. '
         'The Fisher z-transform produces values whose sampling distribution is approximately normal '
         'with variance ≈ 1/(n−3), where n is the number of timepoints. This normalization is a '
         'prerequisite for valid parametric group-level statistics (e.g., t-tests, ANOVAs, regressions) '
         'on edge weights or on graph metrics derived from them.'),
        ('Valid averaging across subjects',
         'Computing group-mean connectivity matrices — or averaging graph metrics derived from weighted '
         'networks — requires that the underlying edge values have symmetric, approximately normal '
         'distributions. Averaging r-values directly introduces bias because the arithmetic mean of '
         'bounded, skewed quantities underestimates the true population correlation, especially when '
         'correlations are strong. With n ≈ 1,000 subjects, even small per-edge biases accumulate into '
         'systematic distortions of the group-mean network and any metrics derived from it. Z-values, '
         'being unbounded and approximately normal, can be legitimately averaged (Silver & Dunlap, 1987).'),
        ('Stabilizing variance across edges',
         'In r-space, the variance of a correlation estimate depends on its magnitude: edges near '
         'r = 0 have higher variance than edges near r = ±1. This heteroscedasticity means that weak '
         'and strong connections contribute unequally to group-level statistics. The Fisher z-transform '
         'stabilizes this variance, ensuring that all edges contribute on an equal statistical footing '
         'to downstream analyses. This is particularly important for weighted graph metrics, where edge '
         'weights directly influence quantities like node strength, weighted clustering, and weighted '
         'efficiency.'),
        ('Compatibility with group-level statistical models',
         'Standard neuroimaging group-analysis frameworks (e.g., Network-Based Statistic, connectome-wide '
         'association studies, mixed-effects models) assume approximately normal residuals. Feeding '
         'z-transformed values into these frameworks satisfies their distributional assumptions. '
         'At n ≈ 1,000, even subtle violations can produce misleading p-values.'),
        ('Effect on graph metrics: monotonic invariance for topology, meaningful impact on weights',
         'Because arctanh(r) is a monotonic function, the rank ordering of edges is preserved. '
         'Proportional thresholding, which retains the top k% of edges by strength, selects exactly the '
         'same edges in r-space and z-space. Therefore all binary/topological graph metrics (binary '
         'clustering, binary efficiency, modularity on the binarized network, degree, small-worldness) '
         'are completely unaffected by the transform. For weighted metrics, the z-transform does change '
         'the weight magnitudes — the nonlinear stretching of high-r values increases the dynamic range '
         'of strong edges. This is appropriate for group analyses because it corrects the statistical '
         'bias described above and makes weighted metrics (node strength, weighted clustering, weighted '
         'efficiency) more suitable for parametric group comparisons.'),
    ]

    for name, desc in z_reasons:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        'Literature support: The Fisher z-transformation before group analysis is the recommended standard '
        'in the functional connectivity literature (Silver & Dunlap, 1987; Lowe et al., 1998; Cole et al., '
        '2010). It is the default in major connectivity analysis toolboxes including CONN and nilearn. '
        'Critically, omitting the z-transform when performing group-level comparisons or correlations with '
        'clinical variables (as planned for this n ≈ 1,000 cohort) would violate the distributional '
        'assumptions of the statistical tests and could inflate false-positive or false-negative rates.'
    )

    doc.add_heading('6.3 Matrix Post-Processing', level=2)
    doc.add_paragraph(
        'After the Fisher z-transformation, three post-processing steps ensure the matrix is suitable '
        'for graph analysis:'
    )
    post_steps = [
        ('Diagonal zeroing', 'Self-connections (diagonal elements) are set to zero after the z-transform, '
         'as they correspond to r = 1 (z = ∞) and are meaningless in a graph context.'),
        ('NaN handling', 'Any NaN values (which can arise from constant time series in parcels with '
         'no signal) are replaced with zero.'),
        ('Symmetry enforcement', 'The matrix is explicitly symmetrized via (M + Mᵀ)/2 to handle '
         'any floating-point asymmetries from the correlation computation.'),
    ]
    for name, desc in post_steps:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # 7. STEP 4 — EDGE HANDLING
    # =========================================================================
    doc.add_heading('7. Step 4 — Edge Handling Strategy', level=1)

    doc.add_paragraph(
        'Before thresholding, the connectivity matrix undergoes edge preparation via the '
        'prepare_graph_matrix() function. Three strategies are available:'
    )

    doc.add_heading('7.1 Positive-Only (Default)', level=2)
    doc.add_paragraph(
        'All negative correlations are set to zero (clipped at 0). Only positive edges are retained.'
    )
    doc.add_paragraph(
        'Scientific rationale: This is the most common approach in the graph theory neuroimaging literature. '
        'Negative functional correlations have an ambiguous neurobiological interpretation — they may reflect '
        'true anti-correlated neural processes, or they may be artifacts of preprocessing steps such as '
        'global signal regression. Most graph theory metrics (particularly those based on shortest paths and '
        'clustering) are designed for non-negative weight matrices. The Brain Connectivity Toolbox (BCT), '
        'which this pipeline uses extensively, requires non-negative weights for most functions. Using '
        'positive-only edges avoids introducing methodological confounds and is consistent with the majority '
        'of published graph theory studies of functional brain networks (Rubinov & Sporns, 2010; van den '
        'Heuvel & Sporns, 2013).'
    )

    doc.add_heading('7.2 Absolute Value', level=2)
    doc.add_paragraph(
        'The absolute value of all correlations is taken, converting negative correlations to positive ones. '
        'This approach treats strong negative correlations as strong connections, under the assumption that '
        'both positive and negative correlations reflect meaningful functional relationships.'
    )

    doc.add_heading('7.3 Signed (Preserve All)', level=2)
    doc.add_paragraph(
        'The raw correlation matrix is passed through unchanged. This approach is only appropriate for '
        'metrics that can handle negative weights, which excludes most standard graph metrics. It is '
        'provided for advanced users with specialized analysis needs.'
    )

    doc.add_paragraph(
        'The default edge_type is "positive_only", configured via the EDGE_TYPE variable. This choice '
        'is recorded in the output metadata for reproducibility.'
    )

    doc.add_page_break()

    # =========================================================================
    # 8. STEP 5 — THRESHOLDING
    # =========================================================================
    doc.add_heading('8. Step 5 — Network Thresholding', level=1)

    doc.add_paragraph(
        'Functional connectivity matrices are dense (fully connected after correlation), but real brain '
        'networks are sparse. Thresholding is necessary to remove weak, potentially spurious connections '
        'and to construct a meaningful network topology.'
    )

    doc.add_heading('8.1 Proportional Thresholding (Default)', level=2)
    doc.add_paragraph(
        'The default approach is proportional (density-based) thresholding, implemented via '
        'bct.threshold_proportional(). This retains a fixed proportion of the strongest edges in the '
        'network. For example, a threshold of 0.15 retains the top 15% of edges by weight.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Scientific rationale: ')
    run.bold = True
    p.add_run(
        'Proportional thresholding is preferred over absolute thresholding for between-subject comparisons '
        'because it ensures that all subjects\' networks have the same density (proportion of possible edges). '
        'This is critical because graph metrics are strongly influenced by network density — a denser network '
        'will have higher clustering, shorter path lengths, and higher efficiency purely as a consequence of '
        'having more edges, independent of any neurobiological differences. By matching density across subjects, '
        'proportional thresholding allows observed differences in graph metrics to be attributed to differences '
        'in network topology rather than confounds of density (van den Heuvel et al., 2017). '
        'Absolute thresholding is also available for exploratory analyses.'
    )

    doc.add_heading('8.2 Default Threshold Range', level=2)
    doc.add_paragraph(
        'The pipeline performs a threshold sweep across the following density levels by default:'
    )
    add_code_block(doc, 'THRESHOLDS = [0.05, 0.10, 0.15, 0.20, 0.25, 0.30, 0.35, 0.40]')
    doc.add_paragraph(
        'This range spans from very sparse (5% density) to moderately dense (40% density) networks, '
        'covering the range most commonly used in the literature.'
    )

    doc.add_heading('8.3 Primary Threshold', level=2)
    doc.add_paragraph(
        'A primary threshold of 0.15 (15% density) is designated for the main analysis outputs. This '
        'threshold is used for:'
    )
    items = [
        'Computing small-worldness (which is computationally expensive)',
        'Generating the main subject-level metrics CSV',
        'Creating visualization figures',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'The choice of 15% density balances network sparsity (removing noise) with retaining sufficient '
        'connections for robust metric computation. This is within the commonly used range of 10–30% in '
        'the literature (Garrison et al., 2015).'
    )

    doc.add_heading('8.4 Post-Thresholding Cleanup', level=2)
    doc.add_paragraph(
        'After thresholding, the matrix is again symmetrized and the diagonal is zeroed to ensure '
        'consistency, as the BCT thresholding functions may introduce minor asymmetries due to '
        'floating-point handling of tied values.'
    )

    doc.add_page_break()

    # =========================================================================
    # 9. STEP 6 — GRAPH METRICS
    # =========================================================================
    doc.add_heading('9. Step 6 — Graph Theory Metric Computation', level=1)

    doc.add_paragraph(
        'At each threshold, the pipeline computes a comprehensive suite of graph metrics on both the '
        'weighted and binarized versions of the thresholded network. All metrics are computed using the '
        'Brain Connectivity Toolbox for Python (bctpy), which implements the algorithms described in '
        'Rubinov & Sporns (2010).'
    )

    doc.add_heading('9.1 Weighted Network Metrics', level=2)

    doc.add_paragraph(
        'The following metrics are computed on the weighted (thresholded but non-binarized) network W:'
    )

    weighted_metrics = [
        ('Global Efficiency', 'efficiency_wei(W)',
         'The average inverse shortest path length across all node pairs. Computed on the weighted '
         'network using weight-to-length conversion (see Section 9.3). Reflects the overall capacity '
         'for parallel information transfer. Range: [0, 1], where 1 indicates a fully connected network.'),
        ('Clustering Coefficient', 'clustering_coef_wu(W)',
         'The fraction of triangles around each node, generalized for weighted undirected networks using '
         'the geometric mean formulation (Onnela et al., 2005). The mean across all nodes provides a '
         'global measure of local connectivity. High clustering indicates a network with dense local '
         'neighborhoods.'),
        ('Characteristic Path Length', 'charpath(D)',
         'The average shortest path length between all pairs of nodes, computed on the distance matrix '
         'D derived from weights. Quantifies the typical number of steps (weighted by connection strength) '
         'required to travel between any two nodes. Shorter path lengths indicate a more integrated network.'),
        ('Modularity (Louvain)', 'community_louvain(W)',
         'The degree to which the network can be partitioned into non-overlapping modules (communities) '
         'with dense within-module connections and sparse between-module connections. Computed using the '
         'Louvain algorithm, which optimizes the modularity quality function Q. See Section 9.4 for '
         'details on the multi-run approach.'),
        ('Node Strength', 'strengths_und(W)',
         'The sum of all edge weights connected to each node. The weighted analog of node degree. Mean '
         'strength provides a global measure of overall connectivity.'),
        ('Betweenness Centrality', 'betweenness_wei(D)',
         'The fraction of all shortest paths in the network that pass through each node. Nodes with high '
         'betweenness serve as critical bridges for information flow. Computed on the distance matrix.'),
        ('Participation Coefficient', 'participation_coef(W, ci)',
         'Quantifies the diversity of inter-modular connections for each node, based on the community '
         'structure identified by Louvain. Nodes with high participation coefficient connect evenly across '
         'multiple modules (connector hubs), while nodes with low participation connect primarily within '
         'their own module (provincial hubs).'),
        ('Local Efficiency', 'efficiency_wei(W, local=True)',
         'The efficiency of the local subgraph around each node (its immediate neighbors). Reflects fault '
         'tolerance — how well communication would be maintained if a node were removed.'),
        ('Assortativity', 'assortativity_wei(W)',
         'The correlation between node strengths at either end of an edge. Positive assortativity indicates '
         'that high-strength nodes tend to connect to other high-strength nodes (assortative mixing). Brain '
         'networks typically show moderate assortativity.'),
        ('Transitivity', 'transitivity_wu(W)',
         'A variant of the clustering coefficient that is normalized by the total number of triangles in '
         'the network rather than averaging per-node clustering. It is less biased by low-degree nodes.'),
    ]

    for name, func, desc in weighted_metrics:
        p = doc.add_paragraph()
        run = p.add_run(f'{name} ')
        run.bold = True
        run2 = p.add_run(f'[{func}]')
        run2.italic = True
        run2.font.size = Pt(9)
        p2 = doc.add_paragraph(desc)
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_after = Pt(6)

    doc.add_heading('9.2 Binary Network Metrics', level=2)

    doc.add_paragraph(
        'The thresholded matrix is also binarized (all non-zero weights set to 1) and the following '
        'binary metrics are computed. Binary metrics complement weighted metrics by characterizing '
        'the network topology independent of edge weight magnitudes.'
    )

    binary_metrics = [
        ('Global Efficiency (binary)', 'efficiency_bin(A)',
         'Inverse shortest path length on the unweighted network.'),
        ('Clustering Coefficient (binary)', 'clustering_coef_bu(A)',
         'Fraction of triangles around each node in the binary network.'),
        ('Characteristic Path Length (binary)', 'charpath(distance_bin(A))',
         'Average shortest path length in hops (unweighted).'),
        ('Modularity (binary)', 'community_louvain(A)',
         'Modularity of the binarized network (best of 10 runs).'),
        ('Node Degree', 'degrees_und(A)',
         'The number of connections (edges) each node has. The fundamental measure of node importance.'),
        ('Betweenness Centrality (binary)', 'betweenness_bin(A)',
         'Fraction of shortest paths passing through each node in the unweighted network.'),
        ('Local Efficiency (binary)', 'efficiency_bin(A, local=True)',
         'Local efficiency on the binary subgraph.'),
        ('Assortativity (binary)', 'assortativity_bin(A)',
         'Degree-degree correlation across edges.'),
        ('Transitivity (binary)', 'transitivity_bu(A)',
         'Global transitivity of the binary network.'),
        ('Rich Club Coefficient', 'rich_club_bu(A)',
         'The rich-club coefficient measures whether high-degree nodes are more densely connected '
         'to each other than expected. Computed only for binary networks.'),
    ]

    for name, func, desc in binary_metrics:
        p = doc.add_paragraph()
        run = p.add_run(f'{name} ')
        run.bold = True
        run2 = p.add_run(f'[{func}]')
        run2.italic = True
        run2.font.size = Pt(9)
        p2 = doc.add_paragraph(desc)
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_after = Pt(6)

    doc.add_heading('9.3 Weight-to-Distance Conversion', level=2)
    doc.add_paragraph(
        'For path-based metrics (characteristic path length, betweenness centrality, global efficiency), '
        'connection weights must be converted to connection lengths/distances. Stronger connections should '
        'correspond to shorter distances.'
    )
    doc.add_paragraph(
        'The pipeline uses BCT\'s weight_conversion(W, "lengths") function, which computes '
        'D_ij = 1 / W_ij for all non-zero edges. Zero-weight (absent) connections are set to infinity '
        'in the distance matrix, indicating no direct path. The diagonal is set to zero (zero distance '
        'to self).'
    )
    doc.add_paragraph(
        'Scientific rationale: The reciprocal transform is the standard approach in the functional '
        'connectivity literature (Rubinov & Sporns, 2010) and is the default in BCT. It provides a '
        'monotonic mapping where strong correlations become short distances, consistent with the '
        'interpretation that strongly connected regions can communicate more efficiently.'
    )

    doc.add_heading('9.4 Modularity: Multi-Run Louvain Algorithm', level=2)
    doc.add_paragraph(
        'The Louvain algorithm for community detection is stochastic — different runs may produce '
        'different community partitions. To address this, the pipeline runs the Louvain algorithm 10 '
        'times for each network and selects the partition with the highest modularity Q value.'
    )
    doc.add_paragraph(
        'Scientific rationale: The Louvain algorithm (Blondel et al., 2008) optimizes the modularity '
        'quality function through a greedy agglomerative process that depends on the order of node '
        'processing. Running multiple iterations and selecting the best result is a standard practice '
        'to mitigate this stochasticity. Ten runs provides a reasonable balance between computational '
        'cost and result stability. The resulting community partition is used both for reporting '
        'modularity Q and for computing the participation coefficient.'
    )

    doc.add_heading('9.5 Density Computation', level=2)
    doc.add_paragraph(
        'Network density is computed as the fraction of possible edges that are present:'
    )
    doc.add_paragraph(
        'density = m / [N × (N−1) / 2]'
    )
    doc.add_paragraph(
        'where m is the number of edges (counted from the upper triangle only, as the network is '
        'undirected) and N is the number of nodes. This is recorded for quality control to verify '
        'that proportional thresholding achieves the target density.'
    )

    doc.add_page_break()

    # =========================================================================
    # 10. STEP 7 — SMALL-WORLD ANALYSIS
    # =========================================================================
    doc.add_heading('10. Step 7 — Small-World Analysis', level=1)

    doc.add_paragraph(
        'Small-worldness characterizes a hallmark property of brain networks: high local clustering '
        'combined with short global path lengths, enabling both specialized local processing and '
        'efficient long-range communication (Watts & Strogatz, 1998).'
    )

    doc.add_heading('10.1 Small-Worldness Coefficient (σ)', level=2)
    doc.add_paragraph(
        'The small-worldness coefficient σ (sigma) is defined as:'
    )
    doc.add_paragraph(
        'σ = γ / λ'
    )
    doc.add_paragraph('where:')
    items = [
        'γ (gamma) = C_real / C_random : the normalized clustering coefficient',
        'λ (lambda) = L_real / L_random : the normalized characteristic path length',
        'C_real and L_real are the clustering coefficient and path length of the observed network',
        'C_random and L_random are the average values across an ensemble of randomized null networks',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'A network is considered small-world if σ > 1, typically γ >> 1 (much higher clustering than '
        'random) and λ ≈ 1 (similar path length to random). Brain networks consistently exhibit '
        'small-world properties with σ values typically in the range of 1.5–5.0.'
    )

    doc.add_heading('10.2 Binary Network Approach', level=2)
    doc.add_paragraph(
        'Small-worldness is computed exclusively on the binarized network. This is a deliberate '
        'methodological choice.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Scientific rationale: ')
    run.bold = True
    p.add_run(
        'The original small-world formulation (Watts & Strogatz, 1998) and the σ coefficient '
        '(Humphries & Gurney, 2008) were defined for binary (unweighted) networks. While weighted '
        'extensions exist, the null model for weighted small-worldness is methodologically problematic: '
        'degree-preserving rewiring of binary networks (Maslov & Sneppen, 2002) is well-defined and '
        'preserves the degree sequence exactly, whereas rewiring weighted networks introduces additional '
        'complexities around weight redistribution. Using binary networks with degree-preserving rewiring '
        'is the most theoretically rigorous approach and is consistent with the standard methodology in '
        'the brain network literature (Bassett & Bullmore, 2006; Bullmore & Sporns, 2009).'
    )

    doc.add_heading('10.3 Null Model: Degree-Preserving Randomization', level=2)
    doc.add_paragraph(
        'The null model uses the Maslov–Sneppen rewiring algorithm (bct.randmio_und), which randomly '
        'swaps pairs of edges while preserving the exact degree sequence of every node. This ensures '
        'that the randomized networks have the same degree distribution as the observed network, so '
        'that any differences in clustering and path length can be attributed to the specific wiring '
        'pattern rather than trivial differences in degree distribution.'
    )
    doc.add_paragraph(
        'Each randomization performs 10 iterations of edge swaps (10× the number of edges), which '
        'is sufficient for thorough randomization while remaining computationally tractable.'
    )

    doc.add_heading('10.4 Number of Randomizations', level=2)
    doc.add_paragraph(
        'By default, N_RANDOM = 100 random networks are generated. The mean clustering and mean path '
        'length across these 100 networks are used as the null reference values. If fewer than 10 valid '
        'random networks are successfully generated (e.g., due to disconnected networks), a warning is '
        'issued. If zero valid networks are generated, σ is set to NaN.'
    )
    doc.add_paragraph(
        'The choice of 100 randomizations balances computational cost with statistical stability. '
        'For publication, 1000 randomizations may be preferred, configurable via the --n-random flag.'
    )

    doc.add_heading('10.5 Computational Optimization', level=2)
    doc.add_paragraph(
        'Small-worldness is only computed at the primary threshold (default: 0.15), not at every '
        'threshold in the sweep. This is because the null model generation is the most computationally '
        'expensive step in the pipeline (O(N_RANDOM × N² × iterations) per subject). Computing at a '
        'single, well-chosen threshold provides the key small-worldness characterization while keeping '
        'runtime tractable for large datasets.'
    )

    doc.add_page_break()

    # =========================================================================
    # 11. STEP 8 — THRESHOLD SWEEP & AUC
    # =========================================================================
    doc.add_heading('11. Step 8 — Threshold Sweep & AUC Integration', level=1)

    doc.add_paragraph(
        'A fundamental challenge in graph theory analysis of brain networks is that results can be '
        'sensitive to the choice of threshold. The pipeline addresses this through a systematic '
        'threshold sweep and AUC (Area Under the Curve) integration.'
    )

    doc.add_heading('11.1 Sweep Procedure', level=2)
    doc.add_paragraph(
        'For each subject/session/run, graph metrics are computed at every threshold in the THRESHOLDS '
        'array (default: 0.05 to 0.40 in steps of 0.05). This produces a metric-vs-threshold curve '
        'for each graph metric.'
    )

    doc.add_heading('11.2 AUC Integration', level=2)
    doc.add_paragraph(
        'The area under the metric-vs-threshold curve is computed using the trapezoidal rule '
        '(numpy.trapezoid). This yields a single summary value for each metric that integrates '
        'information across the full range of thresholds.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Scientific rationale: ')
    run.bold = True
    p.add_run(
        'AUC integration was proposed by Achard & Bullmore (2007) as a way to obtain threshold-independent '
        'summary measures of network properties. By integrating across thresholds, the AUC approach: (1) '
        'reduces dependence on any single arbitrary threshold choice, (2) captures information about how '
        'network properties change across density levels, and (3) provides a single value per metric per '
        'subject suitable for standard statistical analyses (e.g., group comparisons, correlations with '
        'clinical variables). AUC metrics are saved alongside single-threshold metrics in the output.'
    )

    doc.add_heading('11.3 Metrics Included in AUC', level=2)
    doc.add_paragraph('The following metrics are integrated via AUC:')
    auc_metrics = [
        'Global Efficiency (weighted)',
        'Clustering Coefficient (weighted, mean)',
        'Modularity (weighted)',
        'Transitivity (weighted)',
        'Small-worldness σ (where computed)',
    ]
    for m in auc_metrics:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 12. STEP 9 — OUTPUT STRUCTURE
    # =========================================================================
    doc.add_heading('12. Step 9 — Output Structure & Data Products', level=1)

    doc.add_paragraph(
        'The pipeline organizes outputs on a per-subject, per-session basis:'
    )

    add_code_block(doc,
        'output_dir/\n'
        '├── sub-001/\n'
        '│   ├── ses-01/\n'
        '│   │   ├── metrics/\n'
        '│   │   │   ├── sub-001_ses-01_run-01_Glasser_metrics.csv\n'
        '│   │   │   ├── sub-001_ses-01_run-01_Glasser_threshold_sweep.csv\n'
        '│   │   │   ├── sub-001_ses-01_run-01_Glasser_auc.csv\n'
        '│   │   │   └── sub-001_ses-01_Glasser_session_avg.csv\n'
        '│   │   ├── figures/\n'
        '│   │   │   ├── sub-001_ses-01_run-01_Glasser_connectivity_raw.png\n'
        '│   │   │   ├── sub-001_ses-01_run-01_Glasser_connectivity.png\n'
        '│   │   │   ├── sub-001_ses-01_run-01_Glasser_community.png\n'
        '│   │   │   ├── sub-001_ses-01_run-01_Glasser_network.png\n'
        '│   │   │   └── sub-001_ses-01_run-01_Glasser_threshold_sweep.png\n'
        '│   │   └── connectivity_matrices/\n'
        '│   │       ├── sub-001_ses-01_run-01_Glasser_conn.npy\n'
        '│   │       └── sub-001_ses-01_run-01_Glasser_conn_raw.npy\n'
        '│   └── ses-02/\n'
        '│       └── ...\n'
        '└── sub-002/\n'
        '    └── ...'
    )

    doc.add_heading('12.1 Metrics CSV Files', level=2)
    doc.add_paragraph(
        'Each *_metrics.csv file contains one row with all computed metrics for that subject/session/run, '
        'including:'
    )
    csv_fields = [
        'Subject, session, run, and parcellation identifiers',
        'Number of parcels and timepoints',
        'Threshold and density values',
        'Edge type used',
        'All weighted metrics (suffixed _w)',
        'All binary metrics (suffixed _b)',
        'Small-worldness measures (σ, γ, λ, C_real, C_rand, L_real, L_rand)',
    ]
    for f in csv_fields:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('12.2 Threshold Sweep CSV Files', level=2)
    doc.add_paragraph(
        'Each *_threshold_sweep.csv contains one row per threshold, with weighted and binary metrics '
        'at each density level. This enables post-hoc threshold sensitivity analysis.'
    )

    doc.add_heading('12.3 AUC CSV Files', level=2)
    doc.add_paragraph(
        'Each *_auc.csv contains the AUC-integrated metric values for that subject/session/run.'
    )

    doc.add_heading('12.4 Session Average CSV Files', level=2)
    doc.add_paragraph(
        'Each *_session_avg.csv contains the mean metric values averaged across all runs within that '
        'session. This provides one value per session, useful for analyses that do not require run-level '
        'granularity. The number of runs averaged is recorded.'
    )

    doc.add_heading('12.5 Connectivity Matrices', level=2)
    doc.add_paragraph(
        'When --save-matrices is specified, both the raw (full r-value) and processed (after edge '
        'handling) connectivity matrices are saved as NumPy .npy files. These can be loaded with '
        'numpy.load() for custom downstream analyses.'
    )

    doc.add_heading('12.6 Figures', level=2)
    doc.add_paragraph('Five types of figures are generated per subject/session/run:')
    fig_types = [
        ('Raw connectivity matrix', 'Heatmap of the full Pearson correlation matrix (before edge handling), '
         'colored with a diverging RdBu colormap to show positive and negative correlations.'),
        ('Processed connectivity matrix', 'Heatmap after edge handling (e.g., positive-only).'),
        ('Community structure', 'Connectivity matrix reordered by Louvain community assignments, with '
         'black lines demarcating community boundaries. Shows the modular organization.'),
        ('Network graph', 'Spring-layout force-directed graph visualization using NetworkX. Nodes are '
         'colored by community and sized by strength. Only the top 30% of edges are shown for clarity.'),
        ('Threshold sweep', 'Line plots of key metrics across thresholds with AUC values annotated. '
         'Allows visual assessment of metric stability across thresholds.'),
    ]
    for name, desc in fig_types:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # 13. STEP 10 — VISUALIZATION DETAILS
    # =========================================================================
    doc.add_heading('13. Step 10 — Visualization', level=1)

    doc.add_paragraph(
        'All figures are saved at 300 DPI in PNG format. The Matplotlib "Agg" backend is used for '
        'non-interactive rendering, making the pipeline compatible with headless servers and HPC clusters. '
        'The seaborn-v0_8-whitegrid style and HUSL color palette are used for consistent, publication-quality '
        'aesthetics.'
    )

    doc.add_heading('13.1 Connectivity Matrix Heatmaps', level=2)
    doc.add_paragraph(
        'Connectivity matrices are displayed using the RdBu_r (reversed Red-Blue) diverging colormap, '
        'with the color scale set to the 95th percentile of absolute values. This prevents outlier values '
        'from compressing the useful color range.'
    )

    doc.add_heading('13.2 Network Graph Visualization', level=2)
    doc.add_paragraph(
        'Network visualizations use a spring-layout algorithm (Fruchterman–Reingold) with a fixed random '
        'seed (42) for reproducibility. For cleaner visualization, only the top 30% of edges by weight '
        'are displayed. Node sizes scale with strength (or degree), and node colors map to community '
        'assignments. Isolated nodes (with no visible edges after the visualization threshold) are removed '
        'from the plot.'
    )

    doc.add_page_break()

    # =========================================================================
    # 14. CONFIGURATION
    # =========================================================================
    doc.add_heading('14. Configuration Parameters & Defaults', level=1)

    config_rows = [
        ('XCP_OUTPUT_DIR', '/projects/aabdulrasul/TAY/GT/data', 'Path to XCP-D output directory'),
        ('OUTPUT_DIR', '/projects/aabdulrasul/TAY/GT/test', 'Path for pipeline outputs'),
        ('USE_CONCATENATED', 'False', 'Use concatenated (True) or per-run (False) files'),
        ('EDGE_TYPE', 'positive_only', 'Edge handling: positive_only, absolute, or signed'),
        ('THRESHOLDS', '[0.05, 0.10, ..., 0.40]', 'Density thresholds for sweep'),
        ('PRIMARY_THRESHOLD', '0.15', 'Primary threshold for main analysis'),
        ('THRESHOLD_TYPE', 'proportional', 'Thresholding method: proportional or absolute'),
        ('COMPUTE_SMALLWORLD', 'True', 'Whether to compute small-worldness'),
        ('N_RANDOM', '100', 'Number of random networks for null model'),
        ('np.random.seed', '23', 'Random seed for reproducibility'),
    ]

    create_table(doc, ['Parameter', 'Default Value', 'Description'], config_rows)

    doc.add_page_break()

    # =========================================================================
    # 15. SUPPORTED PARCELLATIONS
    # =========================================================================
    doc.add_heading('15. Supported Parcellations', level=1)

    doc.add_paragraph(
        'The pipeline supports any parcellation scheme available in the XCP-D outputs. The EXPECTED_PARCELS '
        'dictionary provides known parcel counts for validation:'
    )

    parc_rows = [
        ('4S156Parcels', '156', 'Schaefer 4-network, 156 parcels'),
        ('4S256Parcels', '256', 'Schaefer 4-network, 256 parcels'),
        ('4S356Parcels', '356', 'Schaefer 4-network, 356 parcels'),
        ('4S456Parcels', '456', 'Schaefer 4-network, 456 parcels'),
        ('4S556Parcels', '556', 'Schaefer 4-network, 556 parcels'),
        ('4S656Parcels', '656', 'Schaefer 4-network, 656 parcels'),
        ('4S756Parcels', '756', 'Schaefer 4-network, 756 parcels'),
        ('4S856Parcels', '856', 'Schaefer 4-network, 856 parcels'),
        ('4S956Parcels', '956', 'Schaefer 4-network, 956 parcels'),
        ('4S1056Parcels', '1056', 'Schaefer 4-network, 1056 parcels'),
        ('Glasser', '360', 'HCP Multi-Modal Parcellation (Glasser et al., 2016)'),
        ('Gordon', '333', 'Gordon parcellation (Gordon et al., 2016)'),
        ('HCP', '360', 'HCP parcellation'),
        ('MIDB', '82', 'MIDB parcellation'),
        ('MyersLabonte', '246', 'Myers-Labonté parcellation'),
        ('Tian', '54', 'Tian subcortical atlas (Tian et al., 2020)'),
    ]

    create_table(doc, ['Parcellation', 'N Parcels', 'Description'], parc_rows)

    doc.add_paragraph('')
    doc.add_paragraph(
        'The Schaefer parcellations (4S series) provide a range of spatial resolutions for multi-scale '
        'analysis. Higher parcel counts provide finer spatial granularity but may reduce signal-to-noise '
        'per parcel and increase computational cost.'
    )

    doc.add_page_break()

    # =========================================================================
    # 16. MATHEMATICAL FORMULATIONS
    # =========================================================================
    doc.add_heading('16. Mathematical Formulations', level=1)

    doc.add_paragraph(
        'This section provides the key mathematical definitions for the computed metrics. '
        'Let W be an N×N weighted adjacency matrix and A be the corresponding binary matrix.'
    )

    doc.add_heading('16.1 Pearson Correlation & Fisher Z-Transform', level=2)
    doc.add_paragraph(
        'For time series x_i and x_j of parcels i and j, each with T timepoints:'
    )
    doc.add_paragraph(
        'r_ij = Σ_t [(x_i(t) − x̄_i)(x_j(t) − x̄_j)] / '
        '[√(Σ_t (x_i(t) − x̄_i)²) × √(Σ_t (x_j(t) − x̄_j)²)]'
    )
    doc.add_paragraph(
        'The Fisher z-transformation is then applied:'
    )
    doc.add_paragraph(
        'z_ij = arctanh(r_ij) = ½ × ln[(1 + r_ij) / (1 − r_ij)]'
    )
    doc.add_paragraph(
        'The sampling distribution of z has approximate variance 1/(n−3), making it suitable for '
        'parametric group-level inference.'
    )

    doc.add_heading('16.2 Proportional Thresholding', level=2)
    doc.add_paragraph(
        'Given a target density p, retain only the top p fraction of edges by weight. '
        'The resulting network has density = p × N(N−1)/2 edges.'
    )

    doc.add_heading('16.3 Global Efficiency (weighted)', level=2)
    doc.add_paragraph(
        'E_global = (1 / N(N−1)) × Σ_{i≠j} 1/d_ij^w'
    )
    doc.add_paragraph(
        'where d_ij^w is the weighted shortest path length between nodes i and j, computed on the '
        'distance matrix D where D_ij = 1/W_ij.'
    )

    doc.add_heading('16.4 Clustering Coefficient (weighted, Onnela)', level=2)
    doc.add_paragraph(
        'C_i = (1 / k_i(k_i − 1)) × Σ_{j,h} (ŵ_ij × ŵ_ih × ŵ_jh)^(1/3)'
    )
    doc.add_paragraph(
        'where ŵ_ij = w_ij / max(w) are the normalized weights, and k_i is the degree of node i.'
    )

    doc.add_heading('16.5 Modularity (Q)', level=2)
    doc.add_paragraph(
        'Q = (1 / 2m) × Σ_{ij} [W_ij − (s_i × s_j / 2m)] × δ(c_i, c_j)'
    )
    doc.add_paragraph(
        'where m = (1/2)Σ_ij W_ij is the total weight, s_i = Σ_j W_ij is the strength of node i, '
        'c_i is the community of node i, and δ is the Kronecker delta.'
    )

    doc.add_heading('16.6 Small-Worldness', level=2)
    doc.add_paragraph(
        'σ = γ / λ = (C_real/C_rand) / (L_real/L_rand)'
    )
    doc.add_paragraph(
        'A network is small-world if σ > 1, which requires γ > 1 (higher clustering than random) '
        'and λ ≈ 1 (similar path length to random).'
    )

    doc.add_heading('16.7 Participation Coefficient', level=2)
    doc.add_paragraph(
        'P_i = 1 − Σ_m (s_i(m) / s_i)²'
    )
    doc.add_paragraph(
        'where s_i(m) is the strength of connections from node i to nodes in module m, and s_i is the '
        'total strength of node i. P_i = 0 if all connections are within one module; P_i → 1 if '
        'connections are distributed evenly across all modules.'
    )

    doc.add_heading('16.8 AUC Integration', level=2)
    doc.add_paragraph(
        'AUC = ∫ f(t) dt ≈ Σ_k [(f(t_k) + f(t_{k+1})) / 2] × (t_{k+1} − t_k)'
    )
    doc.add_paragraph(
        'where f(t) is the metric value at threshold t, computed via the trapezoidal rule.'
    )

    doc.add_page_break()

    # =========================================================================
    # 17. SOFTWARE DEPENDENCIES
    # =========================================================================
    doc.add_heading('17. Software Dependencies & Versions', level=1)

    dep_rows = [
        ('Python', '≥ 3.8', 'Programming language'),
        ('NumPy', '—', 'Array computation and linear algebra'),
        ('Pandas', '—', 'Data manipulation and CSV I/O'),
        ('nibabel', '—', 'Reading CIFTI/NIfTI neuroimaging files'),
        ('bctpy (bct)', '—', 'Brain Connectivity Toolbox for Python — graph metrics'),
        ('SciPy', '—', 'Scientific computing (statistics, integration)'),
        ('matplotlib', '—', 'Figure generation (Agg backend for headless rendering)'),
        ('seaborn', '—', 'Statistical visualization and styling'),
        ('nilearn', '—', 'ConnectivityMeasure for partial correlation (optional)'),
        ('NetworkX', '—', 'Network graph construction and visualization'),
        ('tqdm', '—', 'Progress bars for batch processing'),
        ('argparse', 'stdlib', 'Command-line argument parsing'),
    ]

    create_table(doc, ['Package', 'Version', 'Purpose'], dep_rows)

    doc.add_paragraph('')
    doc.add_paragraph(
        'The bctpy package is a Python implementation of the Brain Connectivity Toolbox originally '
        'developed in MATLAB (Rubinov & Sporns, 2010). It is the de facto standard for graph theory '
        'analysis of brain networks.'
    )

    doc.add_page_break()

    # =========================================================================
    # 18. USAGE INSTRUCTIONS
    # =========================================================================
    doc.add_heading('18. Usage Instructions', level=1)

    doc.add_heading('18.1 Command-Line Interface', level=2)
    doc.add_paragraph('The pipeline is invoked from the command line:')

    add_code_block(doc,
        '# Process a single parcellation:\n'
        'python run_gt.py --parcellation Glasser\n\n'
        '# Process multiple parcellations:\n'
        'python run_gt.py --parcellation Glasser Gordon 4S456Parcels\n\n'
        '# Process all available parcellations:\n'
        'python run_gt.py --all\n\n'
        '# Skip small-worldness (faster):\n'
        'python run_gt.py --parcellation Glasser --no-smallworld\n\n'
        '# Save connectivity matrices:\n'
        'python run_gt.py --parcellation Glasser --save-matrices\n\n'
        '# Skip figure generation:\n'
        'python run_gt.py --parcellation Glasser --no-figures\n\n'
        '# Custom output directory and more random networks:\n'
        'python run_gt.py --parcellation Glasser --output-dir /path/to/output --n-random 1000'
    )

    doc.add_heading('18.2 Command-Line Arguments', level=2)

    arg_rows = [
        ('--parcellation, -p', 'List of parcellation(s) to process'),
        ('--all, -a', 'Process all available parcellations'),
        ('--no-smallworld', 'Skip small-worldness computation (much faster)'),
        ('--n-random', 'Number of random networks for small-worldness (default: 100)'),
        ('--no-figures', 'Skip figure generation'),
        ('--save-matrices', 'Save connectivity matrices as .npy files'),
        ('--output-dir, -o', 'Output directory (overrides default)'),
    ]

    create_table(doc, ['Argument', 'Description'], arg_rows)

    doc.add_heading('18.3 Modifying Default Parameters', level=2)
    doc.add_paragraph(
        'To change default parameters (e.g., thresholds, edge type, input/output paths), edit the '
        'configuration variables at the top of run_gt.py in the "=== CONFIGURATION ===" section.'
    )

    doc.add_page_break()

    # =========================================================================
    # 19. LIMITATIONS & CONSIDERATIONS
    # =========================================================================
    doc.add_heading('19. Limitations & Considerations', level=1)

    limitations = [
        ('Correlation as connectivity measure',
         'Pearson correlation captures linear, zero-lag relationships. It does not capture nonlinear '
         'dependencies, lagged interactions, or directional (causal) connectivity. Alternative measures '
         'such as partial correlation, mutual information, or transfer entropy may provide complementary '
         'information.'),
        ('Static connectivity assumption',
         'The pipeline computes a single connectivity matrix per run, assuming stationarity. Dynamic '
         'functional connectivity methods (e.g., sliding window, hidden Markov models) may reveal '
         'time-varying network properties.'),
        ('Threshold dependence',
         'Despite the AUC approach, some threshold dependence remains. The pipeline mitigates this '
         'through multi-threshold sweep and AUC integration, but users should verify that results are '
         'robust across thresholds.'),
        ('Parcellation choice',
         'Graph metrics are influenced by the parcellation scheme and spatial resolution. Results from '
         'different parcellations are not directly comparable. Multi-scale analysis using multiple '
         'Schaefer resolutions is recommended.'),
        ('Small-worldness limitations',
         'The σ coefficient has known limitations, including sensitivity to network size and density. '
         'The more recently proposed ω (omega) coefficient (Telesford et al., 2011) may provide a more '
         'robust alternative. The pipeline reports all constituent values (C_real, C_rand, L_real, L_rand) '
         'to enable computation of alternative formulations.'),
        ('Louvain stochasticity',
         'Despite the 10-run best-of approach, Louvain community detection remains stochastic. For '
         'high-stakes analyses, increasing the number of runs or using consensus clustering may be '
         'warranted.'),
        ('Motion artifacts',
         'While XCP-D performs confound regression, residual motion artifacts can affect functional '
         'connectivity estimates, particularly for short-distance connections. Users should verify '
         'adequate motion handling in the XCP-D preprocessing step and may wish to include motion '
         'parameters (e.g., mean framewise displacement) as covariates in group-level analyses.'),
        ('Disconnected networks',
         'At very sparse thresholds (e.g., 5% density), networks may become disconnected (contain '
         'isolated components). The pipeline handles this by using include_infinite=False for characteristic '
         'path length computation (only considering connected node pairs) and setting small-worldness to '
         'NaN for disconnected networks. Users should verify network connectedness at their chosen threshold.'),
        ('Computational cost',
         'Small-worldness computation is the primary bottleneck, scaling as O(N_random × N² × rewiring_iterations) '
         'per subject. For large parcellations (e.g., 1000+ parcels) with many subjects, consider using '
         '--no-smallworld for initial exploration and computing small-worldness on a subset or at a single '
         'threshold.'),
    ]

    for name, desc in limitations:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # =========================================================================
    # 20. REFERENCES
    # =========================================================================
    doc.add_heading('20. References', level=1)

    references = [
        'Achard, S., & Bullmore, E. (2007). Efficiency and cost of economical brain functional networks. '
        'PLoS Computational Biology, 3(2), e17.',

        'Bassett, D. S., & Bullmore, E. (2006). Small-world brain networks. The Neuroscientist, 12(6), 512–523.',

        'Blondel, V. D., Guillaume, J.-L., Lambiotte, R., & Lefebvre, E. (2008). Fast unfolding of '
        'communities in large networks. Journal of Statistical Mechanics: Theory and Experiment, 2008(10), P10008.',

        'Bullmore, E., & Sporns, O. (2009). Complex brain networks: graph theoretical analysis of structural '
        'and functional systems. Nature Reviews Neuroscience, 10(3), 186–198.',

        'Garrison, K. A., Scheinost, D., Finn, E. S., Shen, X., & Constable, R. T. (2015). The impact of '
        'rest versus task on network-level connectivity. NeuroImage, 122, 92–104.',

        'Glasser, M. F., Coalson, T. S., Robinson, E. C., Hacker, C. D., Harwell, J., Yacoub, E., ... & '
        'Van Essen, D. C. (2016). A multi-modal parcellation of human cerebral cortex. Nature, 536(7615), 171–178.',

        'Gordon, E. M., Laumann, T. O., Adeyemo, B., Huckins, J. F., Kelley, W. M., & Petersen, S. E. (2016). '
        'Generation and evaluation of a cortical area parcellation from resting-state correlations. Cerebral Cortex, '
        '26(1), 288–303.',

        'Humphries, M. D., & Gurney, K. (2008). Network "small-world-ness": a quantitative method for determining '
        'canonical network equivalence. PLoS ONE, 3(4), e0002051.',

        'Maslov, S., & Sneppen, K. (2002). Specificity and stability in topology of protein networks. Science, '
        '296(5569), 910–913.',

        'Onnela, J.-P., Saramäki, J., Kertész, J., & Kaski, K. (2005). Intensity and coherence of motifs in '
        'weighted complex networks. Physical Review E, 71(6), 065103.',

        'Rubinov, M., & Sporns, O. (2010). Complex network measures of brain connectivity: uses and interpretations. '
        'NeuroImage, 52(3), 1059–1069.',

        'Telesford, Q. K., Joyce, K. E., Hayasaka, S., Burdette, J. H., & Laurienti, P. J. (2011). The ubiquity '
        'of small-world networks. Brain Connectivity, 1(5), 367–375.',

        'Tian, Y., Margulies, D. S., Breakspear, M., & Zalesky, A. (2020). Topographic organization of the human '
        'subcortex unveiled with functional connectivity gradients. Nature Neuroscience, 23(11), 1421–1432.',

        'Van Dijk, K. R. A., Hedden, T., Venkataraman, A., Evans, K. C., Lazar, S. W., & Buckner, R. L. (2010). '
        'Intrinsic functional connectivity as a tool for human connectomics: theory, properties, and optimization. '
        'Journal of Neurophysiology, 103(1), 297–321.',

        'van den Heuvel, M. P., & Sporns, O. (2013). Network hubs in the human brain. Trends in Cognitive Sciences, '
        '17(12), 683–696.',

        'van den Heuvel, M. P., de Lange, S. C., Zalesky, A., Seguin, C., Yeo, B. T. T., & Schmidt, R. (2017). '
        'Proportional thresholding in resting-state fMRI functional connectivity networks and consequences for '
        'patient-control connectome studies: Issues and recommendations. NeuroImage, 152, 437–449.',

        'Watts, D. J., & Strogatz, S. H. (1998). Collective dynamics of "small-world" networks. Nature, '
        '393(6684), 440–442.',

        'Silver, N. C., & Dunlap, W. P. (1987). Averaging correlation coefficients: Should Fisher\'s '
        'z transformation be used? Journal of Applied Psychology, 72(1), 146–148.',

        'Lowe, M. J., Mock, B. J., & Sorenson, J. A. (1998). Functional connectivity in single and '
        'multislice echoplanar imaging using resting-state fluctuations. NeuroImage, 7(2), 119–132.',

        'Cole, D. M., Smith, S. M., & Beckmann, C. F. (2010). Advances and pitfalls in the analysis '
        'and interpretation of resting-state FMRI data. Frontiers in Systems Neuroscience, 4, 8.',
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)

    # =========================================================================
    # SAVE
    # =========================================================================
    output_path = '/home/hassan/Documents/Kimel/PhD_1/Graph_Theory_Pipeline_Documentation.docx'
    doc.save(output_path)
    print(f"Documentation saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    build_document()
